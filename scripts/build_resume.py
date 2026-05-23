"""Build an ATS-friendly fresher resume for Anveet Pal as a .docx file.

ATS rules respected:
- Standard font (Calibri).
- No tables for layout, no text boxes, no header/footer, no images.
- Standard section headings (Summary, Skills, Experience, Projects, Education,
  Certifications) so applicant tracking systems parse them correctly.
- Plain bullet characters; no fancy unicode that ATS may mangle.
- Single column.
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = Path.home() / "Downloads" / "Anveet-Pal-Resume.docx"

DARK = RGBColor(0x14, 0x14, 0x14)
MID = RGBColor(0x33, 0x33, 0x33)


def set_margins(doc: Document, inches: float = 0.6) -> None:
    for section in doc.sections:
        section.top_margin = Inches(inches)
        section.bottom_margin = Inches(inches)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def set_default_font(doc: Document, name: str = "Calibri", size: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = MID
    # Set East Asian font too, so Word doesn't substitute.
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)


def add_bottom_border(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "888888")
    pbdr.append(bottom)
    p_pr.append(pbdr)


def section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = DARK
    add_bottom_border(p)


def line(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: float | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)


def role_line(doc: Document, role: str, company: str, period: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    a = p.add_run(role)
    a.bold = True
    a.font.size = Pt(11.5)
    a.font.color.rgb = DARK
    b = p.add_run(f"  ·  {company}")
    b.font.size = Pt(11.5)
    b.font.color.rgb = MID
    # Right-aligned period via tab stop.
    p.paragraph_format.tab_stops.add_tab_stop(Inches(7.3), WD_ALIGN_PARAGRAPH.RIGHT)
    t = p.add_run("\t" + period)
    t.italic = True
    t.font.color.rgb = MID


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(11)


def main() -> None:
    doc = Document()
    set_margins(doc, 0.55)
    set_default_font(doc, "Calibri", 11)

    # === Header ===
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Anveet Pal")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Computer Science Engineer  ·  Python Developer")
    r.font.size = Pt(11.5)
    r.font.color.rgb = MID

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(
        "8602875539  |  anveetpal12@gmail.com  |  "
        "linkedin.com/in/anveet-pal  |  github.com/anveetpal01"
    )
    r.font.size = Pt(10.5)
    r.font.color.rgb = MID

    # === Summary ===
    section_heading(doc, "Summary")
    p = doc.add_paragraph(
        "Recent B.Tech Computer Science graduate and current Python Developer "
        "with hands-on internship experience in backend scripting, REST APIs, "
        "machine learning, and computer vision. Quick learner who ships clean, "
        "tested code in startup settings — eager to grow with a strong "
        "engineering team."
    )
    p.paragraph_format.space_after = Pt(2)

    # === Skills ===
    section_heading(doc, "Skills")
    skills_groups = [
        ("Languages", "Python, JavaScript, C#, SQL, HTML/CSS"),
        ("Backend & Frameworks", "FastAPI, Flask, Django, .NET, REST APIs"),
        ("Frontend", "React.js, Streamlit"),
        ("Data & ML", "Pandas, Seaborn, XGBoost, scikit-learn, Power BI"),
        ("Computer Vision", "OpenCV, YOLO (Ultralytics), ByteTrack, GPU inference (FP16)"),
        ("Databases & Tools", "MySQL, MSSQL, Git & GitHub, Postman, VS Code, PyCharm"),
        ("Other", "Automation (Selenium), CI/CD basics, ETL"),
    ]
    for label, items in skills_groups:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        a = p.add_run(f"{label}:  ")
        a.bold = True
        a.font.color.rgb = DARK
        a.font.size = Pt(11)
        b = p.add_run(items)
        b.font.size = Pt(11)

    # === Experience ===
    section_heading(doc, "Experience")

    role_line(doc, "Python Developer", "Miraigate Solutions", "Feb 2026 – Present")
    for b in [
        "Build and maintain Python automation scripts and computer-vision "
        "modules across multiple ongoing projects in a fast-paced startup setting.",
        "Collaborate on coding, debugging, testing and project coordination to "
        "support simultaneous deliverables.",
        "Contribute to 4+ concurrent projects, adapting quickly to shifting "
        "priorities and evolving tech requirements.",
    ]:
        bullet(doc, b)

    role_line(doc, "Junior Consultant Intern", "Hosho Digital", "Jul 2025 – Nov 2025")
    for b in [
        "Developed Python scripts for backend data pipelines and REST APIs, "
        "managing databases and ensuring smooth frontend integration.",
        "Automated software tasks with Python and .NET, including data modeling "
        "with the Microsoft Power Platform.",
        "Performed data engineering and analysis with Pandas and Seaborn to "
        "deliver custom client solutions efficiently.",
    ]:
        bullet(doc, b)

    role_line(doc, "Python-ML Intern", "Alveofit (Roundwork Technologies)", "Feb 2025 – Mar 2025")
    for b in [
        "Built and deployed a two-stage XGBoost ML pipeline for real-time "
        "asthma severity prediction; monitored model performance through EDA.",
        "Implemented a Flask REST API with robust input validation and error "
        "handling, integrating the model for production use.",
    ]:
        bullet(doc, b)

    # === Projects ===
    section_heading(doc, "Projects")

    role_line(doc, "CCTV Vehicle Counter", "Computer Vision · Python · YOLO11", "2026")
    bullet(doc, "Real-time computer-vision pipeline that detects, tracks and "
                "counts vehicles from live CCTV streams with low-latency, "
                "GPU-accelerated inference (FP16).")
    bullet(doc, "Designed a line-crossing counter with occlusion handling, "
                "ghost-track suppression, and multi-class aggregation across "
                "concurrent video streams (ByteTrack).")

    role_line(doc, "Disease & Asthma Prediction API",
              "Flask · XGBoost · ETL", "2025")
    bullet(doc, "ETL-driven Flask REST API serving an XGBoost model for "
                "real-time disease severity prediction with feature encoding "
                "and robust error handling.")
    bullet(doc, "Source: github.com/anveetpal01/myProjects/tree/main/xgboost_treat_api")

    role_line(doc, "Multi-User Task Management System",
              "Python · FastAPI · SQL · JWT", "2025")
    bullet(doc, "Scalable backend API with automated tests (Pytest) and cloud "
                "deployment on Render.")
    bullet(doc, "Secured with JWT and password hashing; focused on data "
                "privacy and clean error handling.")

    role_line(doc, "Rewards Management API",
              ".NET 8 · C# · MySQL", "2025")
    bullet(doc, "RESTful API for member accounts, reward-points accrual and "
                "coupon redemption, following clean-architecture principles.")
    bullet(doc, "Source: github.com/anveetpal01/HoshoOnlyMembers")

    # === Education ===
    section_heading(doc, "Education")
    role_line(doc, "B.Tech, Computer Science & Engineering",
              "Medicaps University", "2021 – 2025")

    # === Certifications ===
    section_heading(doc, "Certifications")
    for c in [
        "Microsoft Power BI — Advanced Analytics, Data Visualization, Business Intelligence",
        "Alteryx — ETL, Data Analytics, Analytical Problem Solving",
        "Figma — UI/UX, Web Design",
    ]:
        bullet(doc, c)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
